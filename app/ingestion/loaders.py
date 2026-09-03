"""Извлечение текста из PDF/DOCX/XLSX/CSV."""

from __future__ import annotations

import csv
import importlib.util
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.ingestion.pdf_text import extract_pdf_text

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".xlsx", ".csv"}



class ScannedPdfError(RuntimeError):
    """PDF без текстового слоя, а OCR в этой сборке недоступен."""


def ocr_available() -> bool:
    return importlib.util.find_spec("rapidocr_onnxruntime") is not None


def ocr_pdf(pdf_path: Path) -> str:
    """OCR для сканов через RapidOCR."""
    if not ocr_available():
        raise ScannedPdfError(
            "Файл выглядит как скан: в нём нет текстового слоя. "
            "В этой сборке распознавание отключено — обратитесь к разработчику."
        )

    from app.ingestion.ocr_rapid import ocr_pdf as rapid_ocr_pdf

    return rapid_ocr_pdf(pdf_path)


def extract_text_from_pdf(pdf_path: Path) -> str:
    text = extract_pdf_text(pdf_path)
    return text if text else ocr_pdf(pdf_path)


def extract_text_from_docx(docx_path: Path) -> str:
    doc = DocxDocument(str(docx_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _find_header_row(rows: list[tuple]) -> int:
    """Номер строки с названиями колонок: первая с наибольшим числом ячеек."""
    best_index, best_filled = 0, 0
    for index, row in enumerate(rows[:20]):
        filled = sum(1 for cell in row if cell is not None and str(cell).strip())
        if filled > best_filled:
            best_index, best_filled = index, filled
    return best_index


def extract_text_from_xlsx(xlsx_path: Path) -> str:
    """Каждая строка таблицы — отдельный абзац «Колонка: значение».

    Абзацы разделены пустой строкой: чанкер режет по границам абзацев, и
    позиция реестра не разрывается посередине.
    """
    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    parts: list[str] = []

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        header_index = _find_header_row(rows)
        header = [
            str(cell).strip().replace("\n", " ") if cell is not None else ""
            for cell in rows[header_index]
        ]

        for row in rows[header_index + 1:]:
            cells = [str(cell).strip() if cell is not None else "" for cell in row]
            if not any(cells):
                continue
            line = "; ".join(f"{h}: {v}" for h, v in zip(header, cells) if h and v)
            if line:
                parts.append(line)

    return "\n\n".join(parts)


def extract_text_from_csv(csv_path: Path) -> str:
    with csv_path.open(encoding="utf-8-sig", newline="") as f:
        rows = list(csv.reader(f))
    if not rows:
        return ""

    header = rows[0]
    parts = []
    for row in rows[1:]:
        if not any(row):
            continue
        line = ", ".join(f"{h}: {v}" for h, v in zip(header, row) if v)
        if line:
            parts.append(line)
    return "\n\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix == ".xlsx":
        return extract_text_from_xlsx(path)
    if suffix == ".csv":
        return extract_text_from_csv(path)
    raise ValueError(f"Неподдерживаемый формат файла: {path}")
